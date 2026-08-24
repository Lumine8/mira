"""Turn a mira paper (markdown on disk) into humanized downloadable files:
plain Markdown, Word (.docx), and PDF (.pdf).

All three come from the same markdown, so the download matches what she wrote —
a readable paper with its title, byline, headings, lists, quotes, and the
references/sources she cited. Pure and dependency-light: python-docx for Word,
reportlab for PDF (DejaVu fonts when present, Vera as a fallback so unicode
still renders).

Every public function here is best-effort: the download route wraps them, and a
failure just returns a 500 for that one format — never corrupts the document.
"""

import io
import os
import re

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]\n]+\]\([^)\n]+\))")


def _split_paper(content: str) -> tuple[str, str, str]:
    """Pull the leading ``# Title`` and ``*by …*`` byline off a paper, leaving
    (title, byline, body). Downloads reuse them as a proper title/byline rather
    than rendering them as ordinary body lines."""
    lines = content.replace("\r\n", "\n").split("\n")
    title, byline = "", ""
    rest = lines
    if lines and re.match(r"^#\s+", lines[0]):
        title = re.sub(r"^#\s+", "", lines[0]).strip()
        rest = lines[1:]
    while rest and not rest[0].strip():
        rest = rest[1:]
    if rest and re.match(r"^\*[^*]+\*$", rest[0].strip()):
        byline = rest[0].strip()[1:-1]
        rest = rest[1:]
    return title, byline, "\n".join(rest)


def _inline(text: str) -> list[tuple[str, object]]:
    """Split one line of markdown into inline segments:
    (kind, value) with kind in {text, bold, italic, code, link}."""
    out: list[tuple[str, object]] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append(("text", text[pos : m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            out.append(("bold", tok[2:-2]))
        elif tok.startswith("`"):
            out.append(("code", tok[1:-1]))
        elif tok.startswith("["):
            inner = tok[1:-1]
            label, _, url = inner.partition("](")
            out.append(("link", (label, url)))
        else:
            out.append(("italic", tok[1:-1]))
        pos = m.end()
    if pos < len(text):
        out.append(("text", text[pos:]))
    return out


def _blocks(content: str) -> list[dict]:
    """Chunk a paper body into renderable blocks (heading/quote/list/rule/para)."""
    lines = content.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        hm = re.match(r"^(#{1,3})\s+(.+)$", s)
        if hm:
            blocks.append({"kind": "heading", "level": len(hm.group(1)), "text": hm.group(2).strip()})
            i += 1
            continue
        if re.match(r"^(-{3,}|_{3,}|\*{3,})$", s):
            blocks.append({"kind": "rule"})
            i += 1
            continue
        if s.startswith(">"):
            quote: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append({"kind": "quote", "text": " ".join(quote)})
            continue
        item_match = re.match(r"^(?P<num>\d+)\.\s+(?P<text>.*)$", s) or re.match(
            r"^[-*]\s+(?P<text>.*)$", s
        )
        if item_match:
            ordered = item_match.groupdict().get("num") is not None
            items: list[str] = []
            while i < n:
                st = lines[i].strip()
                im = re.match(r"^\d+\.\s+(.*)$", st) or re.match(r"^[-*]\s+(.*)$", st)
                if not im:
                    break
                items.append(im.group(1))
                i += 1
            blocks.append({"kind": "list", "ordered": ordered, "items": items})
            continue
        para = [s]
        i += 1
        while i < n:
            st = lines[i].strip()
            if (
                not st
                or re.match(r"^#{1,3}\s", st)
                or re.match(r"^(-{3,}|_{3,}|\*{3,})$", st)
                or st.startswith(">")
                or re.match(r"^\d+\.\s", st)
                or re.match(r"^[-*]\s", st)
            ):
                break
            para.append(st)
            i += 1
        blocks.append({"kind": "para", "text": " ".join(para)})
    return blocks


# --------------------------------------------------------------------------- 
# Markdown (.md) — the paper exactly as she wrote it.
# ---------------------------------------------------------------------------


def markdown_bytes(content: str) -> bytes:
    return content.encode("utf-8")


# --------------------------------------------------------------------------- 
# Word (.docx) — python-docx, headings + runs + lists + real hyperlinks.
# ---------------------------------------------------------------------------


def docx_bytes(content: str) -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    title, byline, body = _split_paper(content)
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    def _hyperlink(paragraph, url: str, text: str) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rpr.append(underline)
        run.append(rpr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _runs(paragraph, text: str) -> None:
        for kind, value in _inline(text):
            if kind == "text":
                paragraph.add_run(str(value))
            elif kind == "bold":
                run = paragraph.add_run(str(value))
                run.bold = True
            elif kind == "italic":
                run = paragraph.add_run(str(value))
                run.italic = True
            elif kind == "code":
                run = paragraph.add_run(str(value))
                run.font.name = "Consolas"
            elif kind == "link":
                label, url = value
                _hyperlink(paragraph, url, label)

    doc.add_heading(title or "Mira paper", level=0)
    if byline:
        p = doc.add_paragraph()
        _runs(p, byline)
        for run in p.runs:
            run.italic = True

    for block in _blocks(body):
        kind = block["kind"]
        if kind == "heading":
            doc.add_heading(block["text"], level=min(block["level"], 3))
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            _runs(p, block["text"])
            for run in p.runs:
                run.italic = True
        elif kind == "list":
            style = "List Number" if block["ordered"] else "List Bullet"
            for item in block["items"]:
                p = doc.add_paragraph(style=style)
                _runs(p, item)
        elif kind == "rule":
            continue
        else:
            p = doc.add_paragraph()
            _runs(p, block["text"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- 
# PDF (.pdf) — reportlab with DejaVu (Vera fallback), a clean paper layout.
# ---------------------------------------------------------------------------

_DEJAVU_DIRS = [
    "/usr/share/fonts/truetype/dejavu",
    "C:/Windows/Fonts",
]


def _fonts() -> tuple[str, str, str, str]:
    """Register an unicode-capable family and return (sans, sans-bold, serif,
    mono) font names. Prefers DejaVu (em dashes, curly quotes, ellipsis); the
    system reportlab is bundled with Vera as a graceful fallback."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    base = {
        "DejaVuSans": "DejaVuSans.ttf",
        "DejaVuSans-Bold": "DejaVuSans-Bold.ttf",
        "DejaVuSans-Oblique": "DejaVuSans-Oblique.ttf",
        "DejaVuSans-BoldOblique": "DejaVuSans-BoldOblique.ttf",
        "DejaVuSerif": "DejaVuSerif.ttf",
        "DejaVuSerif-Bold": "DejaVuSerif-Bold.ttf",
        "DejaVuSansMono": "DejaVuSansMono.ttf",
        "DejaVuSansMono-Bold": "DejaVuSansMono-Bold.ttf",
    }
    found: dict[str, str] = {}
    for d in _DEJAVU_DIRS:
        if os.path.isdir(d):
            for name, file in base.items():
                path = os.path.join(d, file)
                if os.path.isfile(path) and name not in found:
                    found[name] = path
    if "DejaVuSans" in found:
        try:
            for name, path in found.items():
                pdfmetrics.registerFont(TTFont(name, path))
            pdfmetrics.registerFontFamily(
                "DejaVuSans",
                normal="DejaVuSans",
                bold=found.get("DejaVuSans-Bold", "DejaVuSans"),
                italic=found.get("DejaVuSans-Oblique", "DejaVuSans"),
                boldItalic=found.get("DejaVuSans-BoldOblique", "DejaVuSans"),
            )
            pdfmetrics.registerFontFamily(
                "DejaVuSerif",
                normal="DejaVuSerif",
                bold=found.get("DejaVuSerif-Bold", "DejaVuSerif"),
                italic="DejaVuSerif",
                boldItalic=found.get("DejaVuSerif-Bold", "DejaVuSerif"),
            )
            return (
                "DejaVuSans",
                "DejaVuSans-Bold",
                "DejaVuSerif",
                found.get("DejaVuSansMono", "DejaVuSans"),
            )
        except Exception:
            pass  # fall through to Vera
    try:
        pdfmetrics.registerFont(TTFont("Vera", "Vera.ttf"))
        pdfmetrics.registerFont(TTFont("VeraBd", "VeraBd.ttf"))
        pdfmetrics.registerFont(TTFont("VeraIt", "VeraIt.ttf"))
        pdfmetrics.registerFontFamily(
            "Vera",
            normal="Vera",
            bold="VeraBd",
            italic="VeraIt",
            boldItalic="VeraBd",
        )
        return "Vera", "VeraBd", "Vera", "Vera"
    except Exception:
        pass
    # Last resort: reportlab's built-in Helvetica (no unicode, but won't crash)
    return "Helvetica", "Helvetica-Bold", "Times-Roman", "Courier"


_ESCAPE = str.maketrans({"&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;"})


def _to_reportlab(text: str, code_font: str) -> str:
    parts: list[str] = []
    for kind, value in _inline(text):
        if kind == "text":
            parts.append(str(value).translate(_ESCAPE))
        elif kind == "bold":
            parts.append(f"<b>{str(value).translate(_ESCAPE)}</b>")
        elif kind == "italic":
            parts.append(f"<i>{str(value).translate(_ESCAPE)}</i>")
        elif kind == "code":
            parts.append(f'<font face="{code_font}">{str(value).translate(_ESCAPE)}</font>')
        elif kind == "link":
            label, url = value
            parts.append(f'<link href="{url.translate(_ESCAPE)}">{label.translate(_ESCAPE)}</link>')
    return "".join(parts)


def pdf_bytes(content: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    title, byline, body = _split_paper(content)
    sans, sans_bold, _serif, mono = _fonts()
    code_font = mono

    def style(name, **kw):
        return ParagraphStyle(name, **kw)

    s_title = style("title", fontName=sans_bold, fontSize=22, leading=27, spaceAfter=4)
    s_byline = style(
        "byline", fontName=sans, fontSize=10, leading=13, textColor=colors.HexColor("#666666"), spaceAfter=18
    )
    s_h = {
        1: style("h1", fontName=sans_bold, fontSize=16, leading=20, spaceBefore=16, spaceAfter=7),
        2: style("h2", fontName=sans_bold, fontSize=13, leading=17, spaceBefore=13, spaceAfter=6),
        3: style("h3", fontName=sans_bold, fontSize=11.5, leading=15, spaceBefore=11, spaceAfter=5),
    }
    s_body = style("body", fontName=sans, fontSize=10.5, leading=15.5, spaceAfter=9, alignment=4)
    s_quote = style(
        "quote",
        fontName=sans,
        fontSize=10.5,
        leading=15.5,
        leftIndent=18,
        rightIndent=10,
        textColor=colors.HexColor("#444444"),
        spaceAfter=9,
    )
    s_item = style("item", fontName=sans, fontSize=10.5, leading=15, spaceAfter=3)

    story: list = []
    if title:
        story.append(Paragraph(_to_reportlab(title, code_font), s_title))
    if byline:
        story.append(Paragraph(_to_reportlab(byline, code_font), s_byline))
    for block in _blocks(body):
        kind = block["kind"]
        if kind == "heading":
            story.append(Paragraph(_to_reportlab(block["text"], code_font), s_h[min(block["level"], 3)]))
        elif kind == "quote":
            story.append(Paragraph(_to_reportlab(block["text"], code_font), s_quote))
        elif kind == "list":
            items = [
                ListItem(Paragraph(_to_reportlab(item, code_font), s_item), leftIndent=14)
                for item in block["items"]
            ]
            story.append(
                ListFlowable(
                    items,
                    bulletType="1" if block["ordered"] else "bullet",
                    start=1 if block["ordered"] else None,
                    bulletFontName=sans,
                    bulletFontSize=10.5,
                    leftIndent=14,
                    spaceAfter=8,
                )
            )
        elif kind == "rule":
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(_to_reportlab(block["text"], code_font), s_body))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=title or "Mira paper",
        author="Mira",
    )
    doc.build(story)
    return buf.getvalue()